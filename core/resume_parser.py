"""
Extracts raw text from an uploaded resume file (.pdf, .docx, or .txt).
Kept deliberately simple: we don't try to understand structure here.
The LLM in tailor.py is responsible for turning raw text into structured
sections. This module's only job is: file bytes -> plain text.
"""

import io


def extract_text(file_bytes: bytes, filename: str) -> str:
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return _extract_docx(file_bytes)
    elif filename_lower.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(
            f"Unsupported file type for '{filename}'. Please upload a .pdf, .docx, or .txt resume."
        )


def _extract_pdf(file_bytes: bytes) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_docx(file_bytes: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also pull text out of any tables (some resumes use table layouts)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)
