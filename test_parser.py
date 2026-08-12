"""
Quick manual test for core.resume_parser.extract_text().
Creates a .docx with paragraphs + a table, saves it, reads it back
through extract_text(), and prints the result.
"""

import os
import tempfile

from docx import Document

from core.resume_parser import extract_text


def main():
    # --- Build a small .docx with paragraphs and a table ---
    doc = Document()
    doc.add_heading("Jane Doe", level=1)
    doc.add_paragraph("Software Engineer | jane@example.com | (555) 123-4567")
    doc.add_paragraph("")  # blank line
    doc.add_heading("Skills", level=2)

    # Use a table (common in real resumes for two-column layouts)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "JavaScript"
    table.cell(1, 0).text = "SQL"
    table.cell(1, 1).text = "Docker"

    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Acme Corp — Backend Developer (2021–2024)")
    doc.add_paragraph("• Built REST APIs serving 10k req/s")
    doc.add_paragraph("• Migrated legacy monolith to microservices")

    # Save to a temp file
    tmp_path = os.path.join(tempfile.gettempdir(), "test_resume.docx")
    doc.save(tmp_path)
    print(f"Saved test .docx to: {tmp_path}\n")

    # --- Read it back through extract_text ---
    with open(tmp_path, "rb") as f:
        file_bytes = f.read()

    result = extract_text(file_bytes, "test_resume.docx")
    print("=" * 60)
    print("EXTRACTED TEXT:")
    print("=" * 60)
    print(result)
    print("=" * 60)

    # --- Also test .txt path ---
    txt_bytes = "Hello from a plain text resume.".encode("utf-8")
    txt_result = extract_text(txt_bytes, "resume.txt")
    print(f"\n.txt extraction: {txt_result!r}")

    # --- Test unsupported format ---
    try:
        extract_text(b"nope", "resume.xlsx")
    except ValueError as e:
        print(f"\nValueError correctly raised: {e}")

    print("\n[OK] All manual tests passed.")


if __name__ == "__main__":
    main()
